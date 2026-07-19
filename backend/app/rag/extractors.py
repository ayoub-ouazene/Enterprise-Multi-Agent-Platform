import re
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.rag.exceptions import KnowledgeExtractionError, KnowledgeValidationError
from app.rag.schemas import ExtractedDocument


SUPPORTED_EXTENSIONS = frozenset({"pdf", "docx", "txt"})
MIME_TYPES = {
    "pdf": frozenset({"application/pdf"}),
    "docx": frozenset(
        {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/zip",
        }
    ),
    "txt": frozenset({"text/plain"}),
}
MAX_PDF_PAGES = 1000
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024


def validate_file_signature(path: Path, extension: str, mime_type: str) -> None:
    extension = extension.lower().lstrip(".")
    normalized_mime = mime_type.lower().split(";", 1)[0].strip()
    if extension not in SUPPORTED_EXTENSIONS:
        raise KnowledgeValidationError("Unsupported knowledge-document file type")
    if normalized_mime not in MIME_TYPES[extension]:
        raise KnowledgeValidationError("File extension and MIME type do not match")

    with path.open("rb") as source:
        prefix = source.read(8192)
    if extension == "pdf" and not prefix.startswith(b"%PDF-"):
        raise KnowledgeValidationError("File extension and content do not match")
    if extension == "docx":
        if not prefix.startswith(b"PK") or not zipfile.is_zipfile(path):
            raise KnowledgeValidationError("File extension and content do not match")
        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
                total_size = sum(item.file_size for item in archive.infolist())
                if (
                    "[Content_Types].xml" not in names
                    or "word/document.xml" not in names
                    or total_size > MAX_DOCX_UNCOMPRESSED_BYTES
                ):
                    raise KnowledgeValidationError("Invalid or unsafe DOCX file")
        except (OSError, zipfile.BadZipFile) as exc:
            raise KnowledgeValidationError("Invalid or unsafe DOCX file") from exc
    if extension == "txt":
        if b"\x00" in prefix:
            raise KnowledgeValidationError("TXT file appears to contain binary data")
        try:
            prefix.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise KnowledgeValidationError("TXT file must use UTF-8 encoding") from exc


def extract_document(path: Path, extension: str) -> ExtractedDocument:
    try:
        if extension == "pdf":
            return _extract_pdf(path)
        if extension == "docx":
            return _extract_docx(path)
        if extension == "txt":
            return _extract_txt(path)
    except KnowledgeExtractionError:
        raise
    except Exception as exc:
        raise KnowledgeExtractionError("Document text could not be extracted") from exc
    raise KnowledgeExtractionError("Unsupported knowledge-document file type")


def _extract_pdf(path: Path) -> ExtractedDocument:
    reader = PdfReader(path, strict=False)
    if len(reader.pages) > MAX_PDF_PAGES:
        raise KnowledgeExtractionError("PDF contains too many pages")
    pages: list[str] = []
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {number}]\n{text.strip()}")
    if not pages:
        raise KnowledgeExtractionError(
            "No usable embedded text was found; scanned PDFs require OCR"
        )
    return ExtractedDocument(text="\n\n".join(pages), source_metadata={"pages": len(reader.pages)})


def _extract_docx(path: Path) -> ExtractedDocument:
    document = DocxDocument(path)
    sections: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            sections.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                sections.append(" | ".join(cells))
    if not sections:
        raise KnowledgeExtractionError("No usable text was found in the DOCX file")
    return ExtractedDocument(
        text="\n\n".join(sections),
        source_metadata={"paragraphs": len(document.paragraphs), "tables": len(document.tables)},
    )


def _extract_txt(path: Path) -> ExtractedDocument:
    try:
        text = path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise KnowledgeExtractionError("TXT file must use UTF-8 encoding") from exc
    if not text.strip():
        raise KnowledgeExtractionError("The TXT file contains no usable text")
    return ExtractedDocument(text=text)


_SPACES = re.compile(r"[ \t]+")
_BLANK_LINES = re.compile(r"\n{3,}")


def clean_text(value: str) -> str:
    value = value.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [_SPACES.sub(" ", line).strip() for line in value.split("\n")]
    return _BLANK_LINES.sub("\n\n", "\n".join(lines)).strip()
