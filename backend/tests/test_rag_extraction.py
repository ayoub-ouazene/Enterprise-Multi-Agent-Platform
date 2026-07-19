from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

import app.rag.extractors as extractors
from app.rag.exceptions import KnowledgeExtractionError, KnowledgeValidationError


def test_txt_extraction(tmp_path: Path) -> None:
    path = tmp_path / "policy.txt"
    path.write_text("Company policy", encoding="utf-8")
    extractors.validate_file_signature(path, "txt", "text/plain")
    assert extractors.extract_document(path, "txt").text == "Company policy"


def test_docx_extraction(tmp_path: Path) -> None:
    path = tmp_path / "manual.docx"
    document = Document()
    document.add_heading("IT Manual")
    document.add_paragraph("Restart the approved service.")
    document.save(path)
    extractors.validate_file_signature(
        path,
        "docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Restart the approved service" in extractors.extract_document(path, "docx").text


def test_pdf_embedded_text_extraction(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "policy.pdf"
    path.write_bytes(b"%PDF-1.7\n")
    page = SimpleNamespace(extract_text=lambda: "Leave policy")
    monkeypatch.setattr(extractors, "PdfReader", lambda *_args, **_kwargs: SimpleNamespace(pages=[page]))
    extractors.validate_file_signature(path, "pdf", "application/pdf")
    assert "Leave policy" in extractors.extract_document(path, "pdf").text


def test_scanned_pdf_fails_safely(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-1.7\n")
    page = SimpleNamespace(extract_text=lambda: "")
    monkeypatch.setattr(extractors, "PdfReader", lambda *_args, **_kwargs: SimpleNamespace(pages=[page]))
    with pytest.raises(KnowledgeExtractionError, match="OCR"):
        extractors.extract_document(path, "pdf")


def test_mime_mismatch_and_binary_txt_are_rejected(tmp_path: Path) -> None:
    path = tmp_path / "wrong.pdf"
    path.write_bytes(b"not a pdf")
    with pytest.raises(KnowledgeValidationError):
        extractors.validate_file_signature(path, "pdf", "application/pdf")
    binary = tmp_path / "bad.txt"
    binary.write_bytes(b"abc\x00def")
    with pytest.raises(KnowledgeValidationError):
        extractors.validate_file_signature(binary, "txt", "text/plain")
